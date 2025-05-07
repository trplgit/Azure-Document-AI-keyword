from flask import Flask, render_template, request, jsonify, redirect, url_for, session
from dotenv import load_dotenv
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
import os
import requests
import datetime
from functools import wraps
from datetime import timedelta
import re
import io
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import fitz  # PyMuPDF for PDF processing
import tempfile
import json

# Load environment variables
load_dotenv()

# Configuration
SEARCH_SERVICE_NAME = os.getenv("SEARCH_SERVICE_NAME")
SEARCH_INDEX_NAME = os.getenv("SEARCH_INDEX_NAME")
API_KEY = os.getenv("API_KEY")
API_VERSION = os.getenv("API_VERSION")
AZURE_STORAGE_CONNECTION_STRING = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
CONTAINER_NAME = os.getenv("CONTAINER_NAME")
ACCOUNT_KEY = os.getenv("ACCOUNT_KEY")

# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", os.urandom(24))
app.permanent_session_lifetime = timedelta(minutes=30)

# Initialize Azure Blob client
blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
container_client = blob_service_client.get_container_client(CONTAINER_NAME)

@app.before_request
def make_session_permanent():
    session.permanent = True

# Valid credentials
# VALID_CREDENTIALS = {
#     os.getenv("ADMIN_USERNAME", "admin"): os.getenv("ADMIN_PASSWORD", "admin123")
# }

VALID_CREDENTIALS = {
    os.getenv("ADMIN_USERNAME"): os.getenv("ADMIN_PASSWORD"),
    os.getenv("USER1_USERNAME"): os.getenv("USER1_PASSWORD"),
    os.getenv("USER2_USERNAME"): os.getenv("USER2_PASSWORD")
}

# Azure Search endpoint
endpoint = f"https://{SEARCH_SERVICE_NAME}.search.windows.net/indexes/{SEARCH_INDEX_NAME}/docs/search?api-version={API_VERSION}"


def escape_query(query: str) -> str:
    # Escape special characters for Lucene
    lucene_special_chars = r'([\+\-\!\(\)\{\}\[\]\^"~\*\?:\\/])'
    query = re.sub(lucene_special_chars, r'\\\1', query)
    return f'"{query}"'  # Wrap in quotes to treat entire thing as a literal string

# Helper for login required
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'username' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# Highlight keywords inside text
def highlight_keywords(text, keywords):
    """Wrap keywords with <mark> tags using different colors."""
    if not text:
        return ""
    
    # Define colors for highlighting
    colors = [
        'background-color: #FFFF00;',  # Yellow
        'background-color: #40E0D0;',  # Turquoise
        'background-color: #FFC0CB;',  # Pink
        'background-color: #90EE90;',  # Green
        'background-color: #98FB98;',  # Bright Green
        'background-color: #ADD8E6;',  # Blue
        'background-color: #FFB6C1;',  # Red
        'background-color: #EE82EE;'   # Violet
    ]
    
    # Create a mapping of keywords to colors
    keyword_colors = {}
    for i, keyword in enumerate(keywords.split()):
        keyword_colors[keyword.lower()] = colors[i % len(colors)]
    
    # Process each keyword
    for keyword, color in keyword_colors.items():
        pattern = re.compile(re.escape(keyword), re.IGNORECASE)
        text = pattern.sub(lambda m: f'<mark style="{color}">{m.group(0)}</mark>', text)
    
    return text

def highlight_keywords_in_docx(blob_client, keywords):
    """Process DOCX file and highlight keywords with different colors."""
    try:
        # Download the blob content
        blob_data = blob_client.download_blob()
        docx_bytes = blob_data.readall()
        
        # Create a Document object from bytes
        doc = Document(io.BytesIO(docx_bytes))
        
        # Define colors for highlighting
        highlight_colors = [
            WD_COLOR_INDEX.YELLOW,      # Yellow
            WD_COLOR_INDEX.TURQUOISE,   # Turquoise
            WD_COLOR_INDEX.PINK,        # Pink
            WD_COLOR_INDEX.GREEN,       # Green
            WD_COLOR_INDEX.BRIGHT_GREEN,# Bright Green
            WD_COLOR_INDEX.BLUE,        # Blue
            WD_COLOR_INDEX.RED,         # Red
            WD_COLOR_INDEX.VIOLET       # Violet
        ]
        
        # Create a mapping of keywords to colors
        keyword_colors = {}
        for i, keyword in enumerate(keywords.split()):
            keyword_colors[keyword.lower()] = highlight_colors[i % len(highlight_colors)]
        
        # Process each paragraph
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                text = run.text
                modified_text = text
                
                # Process each keyword
                for keyword, color in keyword_colors.items():
                    pattern = re.compile(re.escape(keyword), re.IGNORECASE)
                    if pattern.search(modified_text):
                        # Split text by this keyword
                        parts = pattern.split(modified_text)
                        if len(parts) > 1:
                            # Clear the original run
                            run.text = parts[0]
                            
                            # Add highlighted runs for this keyword
                            for i, part in enumerate(parts[1:]):
                                # Add the keyword with its specific color
                                keyword_run = paragraph.add_run(pattern.findall(modified_text)[i])
                                keyword_run.font.highlight_color = color
                                
                                # Add the next part
                                if i < len(parts) - 2:
                                    paragraph.add_run(part)
                            
                            # Update modified_text for next keyword
                            modified_text = pattern.sub('', modified_text)
        
        # Save the modified document
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output.getvalue()
    except Exception as e:
        app.logger.error(f"Error processing DOCX file: {str(e)}")
        return None

def highlight_keywords_in_pdf(blob_client, keywords):
    """Process PDF file and highlight keywords with different colors."""
    try:
        # Download the blob content
        blob_data = blob_client.download_blob()
        pdf_bytes = blob_data.readall()
        
        # Create a temporary file to work with
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_file.write(pdf_bytes)
            temp_path = temp_file.name
        
        # Open the PDF with PyMuPDF
        doc = fitz.open(temp_path)
        
        # Define colors for highlighting
        highlight_colors = [
            (1, 1, 0),       # Yellow
            (0, 0.8, 0.8),   # Turquoise
            (1, 0.75, 0.8),  # Pink
            (0.56, 0.93, 0.56), # Green
            (0.6, 0.98, 0.6),   # Bright Green
            (0.68, 0.85, 0.9),  # Blue
            (1, 0.71, 0.76),    # Red
            (0.93, 0.51, 0.93)  # Violet
        ]
        
        # Create a mapping of keywords to colors
        keyword_colors = {}
        for i, keyword in enumerate(keywords.split()):
            keyword_colors[keyword.lower()] = highlight_colors[i % len(highlight_colors)]
        
        # Process each page
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Search for each keyword
            for keyword, color in keyword_colors.items():
                text_instances = page.search_for(keyword, quads=True)
                
                # Add highlight for each instance
                for inst in text_instances:
                    annot = page.add_highlight_annot(inst)
                    annot.set_colors(stroke=color)
                    annot.update()
        
        # Save the modified PDF
        output = io.BytesIO()
        doc.save(output)
        doc.close()
        
        # Clean up the temporary file
        os.unlink(temp_path)
        
        # Return the modified PDF bytes
        output.seek(0)
        return output.getvalue()
    
    except Exception as e:
        app.logger.error(f"Error processing PDF file: {str(e)}")
        return None

# Login page
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        if username in VALID_CREDENTIALS and VALID_CREDENTIALS[username] == password:
            session['username'] = username
            return redirect(url_for('index'))
        else:
            return render_template('login.html', error='Invalid username or password')
    
    return render_template('login.html')

# Logout
@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect(url_for('login'))

# Home page
@app.route('/')
@login_required
def index():
    return render_template('index.html')

# Search API
@app.route('/search', methods=['POST'])
@login_required
def search():
    user_query = request.form.get('query')
    
    if not user_query:
        return jsonify({'error': 'No query provided'}), 400

    headers = {
        "Content-Type": "application/json",
        "api-key": API_KEY
    }

    escaped_query = escape_query(user_query)
    payload = {
        "search": escaped_query,
        "searchFields": "content,metadata_storage_name",
        "select": "content,metadata_storage_name,metadata_storage_path",
        "top": 10,
        "queryType": "full",
        "searchMode": "all",
        "filter": f"authorized_users eq '{session['username']}'"
    }

    try:
        print("@@@@@@@@@@@ endpoint @@@@@@@@@@@@@@...", endpoint)
        print("@@@@@@@@@@@ headers @@@@@@@@@@@@@@...", headers)
        print("@@@@@@@@@@@ payload @@@@@@@@@@@@@@...", payload)

        response = requests.post(endpoint, headers=headers, json=payload)
        print("@@@@@@@@@@@ payload @@@@@@@@@@@@@@...", response)
        print(response.text)
        print(json.dumps(payload, indent=2))
        response.raise_for_status()
        results = response.json().get("value", [])
        
        for result in results:
            if 'metadata_storage_path' in result:
                try:
                    blob_name = result['metadata_storage_name']
                    blob_client = container_client.get_blob_client(blob_name)
                    
                    if not blob_client.exists():
                        result['view_url'] = None
                        continue
                    
                    # File type detection
                    ext = blob_name.lower().split('.')[-1] if '.' in blob_name else ''
                    result['file_type'] = {
                        'pdf': 'pdf',
                        'doc': 'word', 'docx': 'word',
                        'xls': 'excel', 'xlsx': 'excel',
                        'ppt': 'powerpoint', 'pptx': 'powerpoint',
                        'txt': 'text',
                        'json': 'text',
                        'csv': 'text'
                    }.get(ext, 'other')
                    
                    # For DOCX files, process and highlight keywords
                    if ext in ['doc', 'docx']:
                        highlighted_docx = highlight_keywords_in_docx(blob_client, user_query)
                        if highlighted_docx:
                            # Upload the highlighted version to a temporary blob
                            temp_blob_name = f"highlighted_{blob_name}"
                            temp_blob_client = container_client.get_blob_client(temp_blob_name)
                            temp_blob_client.upload_blob(highlighted_docx, overwrite=True)
                            
                            # Generate SAS token for the highlighted version
                            sas_token = generate_blob_sas(
                                account_name=blob_service_client.account_name,
                                container_name=CONTAINER_NAME,
                                blob_name=temp_blob_name,
                                account_key=ACCOUNT_KEY,
                                permission=BlobSasPermissions(read=True),
                                expiry=datetime.datetime.utcnow() + datetime.timedelta(hours=1),
                                content_disposition='inline'
                            )
                            result['view_url'] = f"{temp_blob_client.url}?{sas_token}"
                        else:
                            # Fallback to original if highlighting fails
                            sas_token = generate_blob_sas(
                                account_name=blob_service_client.account_name,
                                container_name=CONTAINER_NAME,
                                blob_name=blob_name,
                                account_key=ACCOUNT_KEY,
                                permission=BlobSasPermissions(read=True),
                                expiry=datetime.datetime.utcnow() + datetime.timedelta(hours=1),
                                content_disposition='inline'
                            )
                            result['view_url'] = f"{blob_client.url}?{sas_token}"
                    # For PDF files, process and highlight keywords
                    elif ext == 'pdf':
                        highlighted_pdf = highlight_keywords_in_pdf(blob_client, user_query)
                        if highlighted_pdf:
                            # Upload the highlighted version to a temporary blob
                            temp_blob_name = f"highlighted_{blob_name}"
                            temp_blob_client = container_client.get_blob_client(temp_blob_name)
                            temp_blob_client.upload_blob(highlighted_pdf, overwrite=True)
                            
                            # Generate SAS token for the highlighted version
                            sas_token = generate_blob_sas(
                                account_name=blob_service_client.account_name,
                                container_name=CONTAINER_NAME,
                                blob_name=temp_blob_name,
                                account_key=ACCOUNT_KEY,
                                permission=BlobSasPermissions(read=True),
                                expiry=datetime.datetime.utcnow() + datetime.timedelta(hours=1),
                                content_disposition='inline'
                            )
                            result['view_url'] = f"{temp_blob_client.url}?{sas_token}"
                        else:
                            # Fallback to original if highlighting fails
                            sas_token = generate_blob_sas(
                                account_name=blob_service_client.account_name,
                                container_name=CONTAINER_NAME,
                                blob_name=blob_name,
                                account_key=ACCOUNT_KEY,
                                permission=BlobSasPermissions(read=True),
                                expiry=datetime.datetime.utcnow() + datetime.timedelta(hours=1),
                                content_disposition='inline'
                            )
                            result['view_url'] = f"{blob_client.url}?{sas_token}"
                    else:
                        # For other file types, use original blob
                        sas_token = generate_blob_sas(
                            account_name=blob_service_client.account_name,
                            container_name=CONTAINER_NAME,
                            blob_name=blob_name,
                            account_key=ACCOUNT_KEY,
                            permission=BlobSasPermissions(read=True),
                            expiry=datetime.datetime.utcnow() + datetime.timedelta(hours=1),
                            content_disposition='inline'
                        )
                        result['view_url'] = f"{blob_client.url}?{sas_token}"
                    
                    # Metadata
                    props = blob_client.get_blob_properties()
                    result['file_size'] = props.size
                    result['last_modified'] = props.last_modified.isoformat()
                    
                except Exception as e:
                    app.logger.error(f"Error generating SAS URL for blob {blob_name}: {str(e)}")
                    result['view_url'] = None

            # Highlight content
            if 'content' in result:
                result['highlighted_content'] = highlight_keywords(result['content'], user_query)
            else:
                result['highlighted_content'] = ''

        return jsonify({'results': results})
    
    except requests.exceptions.RequestException as e:
        app.logger.error(f"Search API error: {str(e)}")
        return jsonify({'error': 'Search service error. Please try again later.'}), 500

if __name__ == '__main__':
  app.run(debug=True, port=5001, use_reloader=False)